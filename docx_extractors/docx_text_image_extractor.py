"""
Extract text and images separately from DOCX files.
"""

import os
from docx import Document
from PIL import Image
import io
from describe_image import describe_image
from .docx_extractor_base import DocxExtractor


class DocxTextImageExtractor(DocxExtractor):
    """Extract text and images separately from DOCX."""

    def extract_text(self, doc):
        """Extract text from a DOCX document."""
        text_content = []

        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                # Check if it's a heading
                if para.style.name.startswith("Heading"):
                    level = (
                        int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    )
                    text_content.append(f"{'#' * level} {para.text}")
                else:
                    text_content.append(para.text)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():  # Only add non-empty rows
                    text_content.append(row_text)

        # Join with double newlines to ensure proper paragraph separation
        return "\n\n".join(text_content)

    def extract_images(self, doc):
        """Extract images from a DOCX document."""
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                images.append(image_data)
        return images

    def save_image(self, image_data, output_dir, docx_filename, image_index):
        """Save an image extracted from a DOCX document."""
        image = Image.open(io.BytesIO(image_data))
        image_name = f"{docx_filename}_image_{image_index}.png"
        img_path = os.path.join(output_dir, image_name)
        image.save(img_path, "PNG")
        return img_path

    def extract(self, docx_path: str, output_dir: str) -> str:
        """Process DOCX file by extracting text and images separately."""
        try:
            docx_filename = os.path.splitext(os.path.basename(docx_path))[0]
            doc = Document(docx_path)

            text_content = self.extract_text(doc)
            image_data_list = self.extract_images(doc)

            md_content = f"# {docx_filename}\n\n"
            md_content += text_content + "\n\n"

            for index, image_data in enumerate(image_data_list):
                img_path = self.save_image(
                    image_data, output_dir, docx_filename, index + 1
                )
                image_description = describe_image(img_path)
                md_content += f"[Image {index + 1}]\n"
                md_content += f"Description: {image_description}\n\n"

                # Delete the temporary image file
                os.remove(img_path)

            md_file_path = os.path.join(output_dir, f"{docx_filename}.md")
            with open(md_file_path, "w", encoding="utf-8") as md_file:
                md_file.write(md_content)

            return md_file_path

        except Exception as e:
            print(f"Error processing DOCX: {str(e)}")
            raise
